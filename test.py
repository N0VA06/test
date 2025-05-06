import os
from pathlib import Path
import torch
import logging
import numpy as np
from transformers import AutoTokenizer, AutoModel
from langchain_community.document_loaders import DirectoryLoader, PyPDFLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document  # Import the Document class from LangChain
import pickle
from docx import Document as DocxDocument  # Import for handling .docx files

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Constants
COLLECTION_NAME = "document_collection"
EMBEDDING_DIM = 1024  # multilingual-e5-large-instruct embedding dimension
MODEL_NAME = "intfloat/multilingual-e5-large-instruct"
SAVE_PATH = "./models/multilingual-e5-large-instruct"

# Save model and tokenizer locally
tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
tokenizer.save_pretrained(SAVE_PATH)

model = AutoModel.from_pretrained(MODEL_NAME)
model.save_pretrained(SAVE_PATH)

print(f"Model and tokenizer saved to {SAVE_PATH}")

def try_init_qdrant():
    """Try to initialize Qdrant client with proper error handling"""
    try:
        # Import inside function to handle missing package
        from qdrant_client import QdrantClient
        
        # Try connecting to Qdrant server
        logger.info("Attempting to connect to Qdrant server")
        client = QdrantClient(url="http://localhost:6333")
        
        # Test connection with a simple operation
        try:
            client.get_collections()
            logger.info("Successfully connected to Qdrant server")
            
            # Check if collection exists, create if not
            collections = client.get_collections().collections
            collection_names = [collection.name for collection in collections]
            
            if COLLECTION_NAME not in collection_names:
                # Create collection with proper parameters
                client.create_collection(
                    collection_name=COLLECTION_NAME,
                    vectors_config={
                        "size": EMBEDDING_DIM,
                        "distance": "Cosine"
                    }
                )
                logger.info(f"Collection '{COLLECTION_NAME}' created")
            else:
                logger.info(f"Collection '{COLLECTION_NAME}' already exists")
            
            return client
        except Exception as e:
            logger.error(f"Qdrant server connection failed: {str(e)}")
            return None
            
    except ImportError:
        logger.warning("Qdrant client not installed")
        return None
    except Exception as e:
        logger.error(f"Error initializing Qdrant: {str(e)}")
        return None

# Initialize the embedding model
def init_embedding_model():
    """Initialize the embedding model locally"""
    logger.info(f"Loading embedding model from local directory: {MODEL_NAME}")
    try:
        # Specify the local path where the model and tokenizer are stored
        local_model_path = "./models/multilingual-e5-large-instruct"
        
        tokenizer = AutoTokenizer.from_pretrained(local_model_path)
        model = AutoModel.from_pretrained(local_model_path)
        device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        logger.info(f"Using device: {device}")
        model.to(device)
        return tokenizer, model, device
    except Exception as e:
        logger.error(f"Error loading embedding model locally: {str(e)}")
        raise

def compute_embeddings(texts, tokenizer, model, device):
    """Compute embeddings for a list of texts using multilingual-e5-large-instruct"""
    embeddings = []
    batch_size = 8  # Smaller batch size to avoid memory issues
    
    for i in range(0, len(texts), batch_size):
        batch_texts = texts[i:i+batch_size]
        logger.debug(f"Computing embeddings for batch {i//batch_size + 1}/{(len(texts)-1)//batch_size + 1}")
        
        # Prepare instruction format (recommended for E5 models)
        instruction_texts = [f"passage: {text}" for text in batch_texts]
        
        # Prepare the model inputs
        inputs = tokenizer(
            instruction_texts, 
            padding=True, 
            truncation=True, 
            max_length=512, 
            return_tensors="pt"
        ).to(device)
        
        # Get the embeddings
        with torch.no_grad():
            outputs = model(**inputs)
            # Mean pooling
            attention_mask = inputs["attention_mask"]
            embeddings_batch = mean_pooling(outputs.last_hidden_state, attention_mask)
            embeddings.extend(embeddings_batch.cpu().numpy())
    
    return embeddings

def mean_pooling(token_embeddings, attention_mask):
    """Mean pooling to get sentence embeddings"""
    input_mask_expanded = attention_mask.unsqueeze(-1).expand(token_embeddings.size()).float()
    sum_embeddings = torch.sum(token_embeddings * input_mask_expanded, 1)
    sum_mask = torch.clamp(input_mask_expanded.sum(1), min=1e-9)
    return sum_embeddings / sum_mask

# Load documents from the content directory
def load_documents(content_dir):
    """Load documents from the content directory"""
    documents = []
    content_path = Path(content_dir)
    
    if not content_path.exists():
        logger.error(f"Content directory '{content_dir}' does not exist")
        return documents
    
    # Make sure the directory exists
    content_path.mkdir(parents=True, exist_ok=True)
    
    # PDF loader
    try:
        pdf_loader = DirectoryLoader(
            content_dir, 
            glob="**/*.pdf", 
            loader_cls=PyPDFLoader,
            show_progress=True
        )
        pdf_docs = pdf_loader.load()
        logger.info(f"Loaded {len(pdf_docs)} PDF documents from {content_dir}")
        documents.extend(pdf_docs)
    except Exception as e:
        logger.error(f"Error loading PDF documents: {str(e)}")
    
    # Text loader for .txt files
    try:
        text_loader = DirectoryLoader(
            content_dir, 
            glob="**/*.txt", 
            loader_cls=TextLoader,
            show_progress=True
        )
        text_docs = text_loader.load()
        logger.info(f"Loaded {len(text_docs)} text documents from {content_dir}")
        documents.extend(text_docs)
    except Exception as e:
        logger.error(f"Error loading text documents: {str(e)}")
    
    # Custom loader for .docx files
    try:
        docx_files = list(content_path.rglob("*.docx"))
        for docx_file in docx_files:
            doc = DocxDocument(docx_file)
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            
            # Simulate page metadata
            total_pages = len(paragraphs)
            for page_number, paragraph in enumerate(paragraphs):
                metadata = {
                    "author": doc.core_properties.author,
                    "title": doc.core_properties.title,
                    "subject": doc.core_properties.subject,
                    "keywords": doc.core_properties.keywords,
                    "last_modified_by": doc.core_properties.last_modified_by,
                    "created": doc.core_properties.created.isoformat() if doc.core_properties.created else None,
                    "modified": doc.core_properties.modified.isoformat() if doc.core_properties.modified else None,
                    "source": str(docx_file),
                    "total_pages": total_pages,
                    "page": page_number,
                    "page_label": str(page_number + 1)
                }
                documents.append(Document(page_content=paragraph, metadata=metadata))
        logger.info(f"Loaded {len(docx_files)} DOCX documents from {content_dir}")
    except Exception as e:
        logger.error(f"Error loading DOCX documents: {str(e)}")
    
    logger.info(f"Loaded {len(documents)} total documents from {content_dir}")
    return documents

# Split documents into chunks
def split_documents(documents):
    """Split documents into chunks for better indexing"""
    if not documents:
        logger.warning("No documents to split")
        return []
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
    )
    
    chunks = text_splitter.split_documents(documents)
    logger.info(f"Split into {len(chunks)} chunks")
    return chunks

# Index documents in storage
def index_documents(chunks, storage, tokenizer, model, device, batch_size=100):
    """Index document chunks in storage (Qdrant or local)"""
    if not chunks:
        logger.warning("No chunks to index")
        return
    
    # Check if storage is from Qdrant or local
    try:
        from qdrant_client import QdrantClient
        using_qdrant = isinstance(storage, QdrantClient)
    except ImportError:
        using_qdrant = False
    
    if not using_qdrant and not isinstance(storage, LocalVectorStorage):
        logger.error("Unknown storage type")
        return
    
    for i in range(0, len(chunks), batch_size):
        batch_chunks = chunks[i:i+batch_size]
        texts = [chunk.page_content for chunk in batch_chunks]
        metadatas = [chunk.metadata for chunk in batch_chunks]
        
        # Compute embeddings for batch
        embeddings = compute_embeddings(texts, tokenizer, model, device)
        
        if using_qdrant:
            # Using Qdrant
            from qdrant_client import models
            
            # Prepare points for Qdrant
            points = []
            for j, (text, metadata, embedding) in enumerate(zip(texts, metadatas, embeddings)):
                point_id = i + j
                points.append(models.PointStruct(
                    id=point_id,
                    vector=embedding.tolist(),  # Convert to list for compatibility
                    payload={
                        "text": text,
                        "metadata": metadata
                    }
                ))
            
            # Upload points to Qdrant
            storage.upsert(
                collection_name=COLLECTION_NAME,
                points=points
            )
        else:
            # Using local storage
            payloads = []
            for text, metadata in zip(texts, metadatas):
                payloads.append({
                    "text": text,
                    "metadata": metadata
                })
            
            storage.upsert(embeddings, payloads)
        
        logger.info(f"Indexed batch {i//batch_size + 1}/{(len(chunks)-1)//batch_size + 1}, {len(batch_chunks)} chunks")

# Main function to run the indexing process
def main(content_dir, use_local_storage=False):
    """Main function to run the indexing process"""
    # Initialize Qdrant client or local storage
    storage = try_init_qdrant()
    
    # Initialize embedding model
    tokenizer, model, device = init_embedding_model()
    
    # Load documents
    documents = load_documents(content_dir)
    
    # Split into chunks
    chunks = split_documents(documents)
    
    # Index in storage
    index_documents(chunks, storage, tokenizer, model, device)
    
    logger.info(f"Indexing complete. {len(chunks)} chunks processed.")
    
    return storage

# Search function for testing
def search_documents(query, storage, tokenizer, model, device, top_k=5, keyword=None):
    """Hybrid search: Combine vector similarity with keyword filtering"""
    # Compute query embedding
    query_embedding = compute_embeddings([query], tokenizer, model, device)[0]
    
    # Check if storage is from Qdrant or local
    try:
        from qdrant_client import QdrantClient
        using_qdrant = isinstance(storage, QdrantClient)
    except ImportError:
        using_qdrant = False
    
    # Perform vector similarity search
    if using_qdrant:
        # Using Qdrant
        results = storage.search(
            collection_name=COLLECTION_NAME,
            query_vector=query_embedding.tolist(),
            limit=top_k
        )
        
        # Format results
        formatted_results = []
        for res in results:
            formatted_results.append({
                "text": res.payload["text"],
                "metadata": res.payload["metadata"],
                "score": res.score
            })
    
    # Apply keyword filtering if a keyword is provided
    if keyword:
        filtered_results = [
            result for result in formatted_results
            if keyword.lower() in result["text"].lower() or
               any(keyword.lower() in str(value).lower() for value in result["metadata"].values())
        ]
    else:
        filtered_results = formatted_results
    
    return filtered_results

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Document Indexing with E5 Embeddings")
    parser.add_argument("--content_dir", type=str, default="content", help="Path to the content directory")
    parser.add_argument("--search", type=str, help="Search query to test the system")
    parser.add_argument("--keyword", type=str, help="Keyword for hybrid search")
    args = parser.parse_args()
    
    # Set the content directory path
    content_directory = args.content_dir
    
    # Run indexing
    storage = main(content_directory, use_local_storage=args.local)
    
    # Test search if requested
    if args.search:
        tokenizer, model, device = init_embedding_model()
        results = search_documents(args.search, storage, tokenizer, model, device, keyword=args.keyword)
        
        print("\nSearch Results:")
        print("-" * 80)
        for i, result in enumerate(results):
            print(f"Result {i+1} (Score: {result['score']:.4f}):")
            print(f"Text: {result['text']}")
            print(f"Source: {result['metadata'].get('source', 'Unknown')}")
            print(f"Page Number: {result['metadata'].get('page_label', 'Unknown')}")
            print("-" * 80)