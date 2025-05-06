import os
from pathlib import Path
import torch
import logging
import numpy as np
import requests
import json
from langchain_community.document_loaders import DirectoryLoader, PyPDFLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document  # Import the Document class from LangChain
import pickle
from docx import Document as DocxDocument  # Import for handling .docx files

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Constants
COLLECTION_NAME = "document_collection"
EMBEDDING_DIM = 1024  # Match the dimension of your existing Qdrant collection
EMBEDDING_DEPLOYMENT = "text-embedding-ada-002"  # Replace with your deployment name

# Azure OpenAI Configuration
AZURE_OPENAI_ENDPOINT = "https://embdeing.openai.azure.com/"  # Use your actual endpoint
AZURE_OPENAI_KEY = ""  # Replace with your API key
AZURE_OPENAI_API_VERSION = ""  # Use the appropriate API version

# Load API key and endpoint from environment variables if available
import os
if os.getenv("AZURE_OPENAI_API_KEY"):
    AZURE_OPENAI_KEY = os.getenv("AZURE_OPENAI_API_KEY")
if os.getenv("AZURE_OPENAI_ENDPOINT"):
    AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
if os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT"):
    EMBEDDING_DEPLOYMENT = os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT")

def try_init_qdrant():
    """Try to initialize Qdrant client with proper error handling"""
    try:
        # Import inside function to handle missing package
        from qdrant_client import QdrantClient
        
        # Try connecting to Qdrant server
        logger.info("Attempting to connect to Qdrant server")
        client = QdrantClient(url="http://localhost:6333")
        
        # Test connection with a simple operation
        try:
            client.get_collections()
            logger.info("Successfully connected to Qdrant server")
            
            # Check if collection exists, create if not
            collections = client.get_collections().collections
            collection_names = [collection.name for collection in collections]
            
            if COLLECTION_NAME not in collection_names:
                # Create collection with proper parameters
                client.create_collection(
                    collection_name=COLLECTION_NAME,
                    vectors_config={
                        "size": EMBEDDING_DIM,
                        "distance": "Cosine"
                    }
                )
                logger.info(f"Collection '{COLLECTION_NAME}' created")
            else:
                logger.info(f"Collection '{COLLECTION_NAME}' already exists")
            
            return client
        except Exception as e:
            logger.error(f"Qdrant server connection failed: {str(e)}")
            return None
            
    except ImportError:
        logger.warning("Qdrant client not installed")
        return None
    except Exception as e:
        logger.error(f"Error initializing Qdrant: {str(e)}")
        return None

# Function to get embeddings from Azure OpenAI
def get_azure_embeddings(texts):
    """Get embeddings from Azure OpenAI API"""
    embeddings = []
    batch_size = 16  # Adjust based on Azure service limits
    
    headers = {
        "Content-Type": "application/json",
        "api-key": AZURE_OPENAI_KEY
    }
    
    # Log configuration for debugging
    logger.info(f"Azure OpenAI Configuration:")
    logger.info(f"Endpoint: {AZURE_OPENAI_ENDPOINT}")
    logger.info(f"API Version: {AZURE_OPENAI_API_VERSION}")
    logger.info(f"Deployment: {EMBEDDING_DEPLOYMENT}")
    if AZURE_OPENAI_KEY and len(AZURE_OPENAI_KEY) > 5:
        logger.info(f"API Key (first 5 chars): {AZURE_OPENAI_KEY[:5]}{'*' * 10}")
    else:
        logger.warning("API Key not provided or too short")
    
    for i in range(0, len(texts), batch_size):
        batch_texts = texts[i:i+batch_size]
        logger.debug(f"Computing embeddings for batch {i//batch_size + 1}/{(len(texts)-1)//batch_size + 1}")
        
        # Prepare the request payload
        payload = {
            "input": batch_texts
        }
        
        # Make the request to Azure OpenAI
        # Try multiple URL formats to handle different configurations
        urls_to_try = [
            "https://embdeing.openai.azure.com/openai/deployments/text-embedding-ada-002/embeddings?api-version=2023-05-15",
            f"{AZURE_OPENAI_ENDPOINT}openai/deployments/{EMBEDDING_DEPLOYMENT}/embeddings?api-version={AZURE_OPENAI_API_VERSION}"
        ]
        
        response = None
        successful = False
        
        for url in urls_to_try:
            logger.info(f"Attempting request to URL: {url}")
            try:
                response = requests.post(url, headers=headers, json=payload, timeout=30)
                logger.info(f"Response status code: {response.status_code}")
                
                if response.status_code == 200:
                    successful = True
                    break
                else:
                    logger.error(f"Error response: {response.text}")
                    
                    # Try different API versions if 401/403 error occurs
                    if response.status_code in [401, 403]:
                        for version in SUPPORTED_API_VERSIONS:
                            if version == AZURE_OPENAI_API_VERSION:
                                continue
                            alt_url = url.replace(AZURE_OPENAI_API_VERSION, version)
                            logger.info(f"Trying with different API version: {version}")
                            alt_response = requests.post(alt_url, headers=headers, json=payload, timeout=30)
                            if alt_response.status_code == 200:
                                logger.info(f"Success with API version: {version}")
                                response = alt_response
                                successful = True
                                break
                        
                        if successful:
                            break
            except requests.exceptions.RequestException as e:
                logger.error(f"Request exception for URL {url}: {str(e)}")
        
        if not successful:
            logger.error("All URL attempts failed")
            # Add empty embeddings as fallback
            for _ in range(len(batch_texts)):
                embeddings.append([0.0] * EMBEDDING_DIM)
            continue
            
        try:
            # Extract embeddings from the response
            response_data = response.json()
            
            # Check if we got a proper response structure
            if "data" not in response_data:
                logger.error(f"Unexpected response format: {response_data}")
                # Add empty embeddings as fallback
                for _ in range(len(batch_texts)):
                    embeddings.append([0.0] * EMBEDDING_DIM)
                continue
                
            batch_embeddings = [item["embedding"] for item in response_data["data"]]
            
            # Resize embeddings to match the required dimension
            resized_embeddings = []
            for emb in batch_embeddings:
                if len(emb) > EMBEDDING_DIM:
                    # Truncate to match the collection's dimension
                    resized_embeddings.append(emb[:EMBEDDING_DIM])
                elif len(emb) < EMBEDDING_DIM:
                    # Pad with zeros if needed
                    padded = emb + [0.0] * (EMBEDDING_DIM - len(emb))
                    resized_embeddings.append(padded)
                else:
                    resized_embeddings.append(emb)
            
            embeddings.extend(resized_embeddings)
            
        except Exception as e:
            logger.error(f"Error processing embedding response: {str(e)}")
            # Add empty embeddings as fallback
            for _ in range(len(batch_texts)):
                embeddings.append([0.0] * EMBEDDING_DIM)
    
    return embeddings

# Load documents from the content directory
def load_documents(content_dir):
    """Load documents from the content directory"""
    documents = []
    content_path = Path(content_dir)
    
    if not content_path.exists():
        logger.error(f"Content directory '{content_dir}' does not exist")
        return documents
    
    # Make sure the directory exists
    content_path.mkdir(parents=True, exist_ok=True)
    
    # PDF loader
    try:
        pdf_loader = DirectoryLoader(
            content_dir, 
            glob="**/*.pdf", 
            loader_cls=PyPDFLoader,
            show_progress=True
        )
        pdf_docs = pdf_loader.load()
        logger.info(f"Loaded {len(pdf_docs)} PDF documents from {content_dir}")
        documents.extend(pdf_docs)
    except Exception as e:
        logger.error(f"Error loading PDF documents: {str(e)}")
    
    # Text loader for .txt files
    try:
        text_loader = DirectoryLoader(
            content_dir, 
            glob="**/*.txt", 
            loader_cls=TextLoader,
            show_progress=True
        )
        text_docs = text_loader.load()
        logger.info(f"Loaded {len(text_docs)} text documents from {content_dir}")
        documents.extend(text_docs)
    except Exception as e:
        logger.error(f"Error loading text documents: {str(e)}")
    
    # Custom loader for .docx files
    try:
        docx_files = list(content_path.rglob("*.docx"))
        for docx_file in docx_files:
            doc = DocxDocument(docx_file)
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            
            # Simulate page metadata
            total_pages = len(paragraphs)
            for page_number, paragraph in enumerate(paragraphs):
                metadata = {
                    "author": doc.core_properties.author,
                    "title": doc.core_properties.title,
                    "subject": doc.core_properties.subject,
                    "keywords": doc.core_properties.keywords,
                    "last_modified_by": doc.core_properties.last_modified_by,
                    "created": doc.core_properties.created.isoformat() if doc.core_properties.created else None,
                    "modified": doc.core_properties.modified.isoformat() if doc.core_properties.modified else None,
                    "source": str(docx_file),
                    "total_pages": total_pages,
                    "page": page_number,
                    "page_label": str(page_number + 1)
                }
                documents.append(Document(page_content=paragraph, metadata=metadata))
        logger.info(f"Loaded {len(docx_files)} DOCX documents from {content_dir}")
    except Exception as e:
        logger.error(f"Error loading DOCX documents: {str(e)}")
    
    logger.info(f"Loaded {len(documents)} total documents from {content_dir}")
    return documents

# Split documents into chunks
def split_documents(documents):
    """Split documents into chunks for better indexing"""
    if not documents:
        logger.warning("No documents to split")
        return []
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
    )
    
    chunks = text_splitter.split_documents(documents)
    logger.info(f"Split into {len(chunks)} chunks")
    return chunks

# Index documents in storage
def index_documents(chunks, storage, batch_size=100):
    """Index document chunks in storage (Qdrant or local)"""
    if not chunks:
        logger.warning("No chunks to index")
        return
    
    # Check if storage is from Qdrant or local
    try:
        from qdrant_client import QdrantClient
        using_qdrant = isinstance(storage, QdrantClient)
    except ImportError:
        using_qdrant = False
    
    if not using_qdrant and not isinstance(storage, LocalVectorStorage):
        logger.error("Unknown storage type")
        return
    
    for i in range(0, len(chunks), batch_size):
        batch_chunks = chunks[i:i+batch_size]
        texts = [chunk.page_content for chunk in batch_chunks]
        metadatas = [chunk.metadata for chunk in batch_chunks]
        
        # Compute embeddings for batch using Azure OpenAI
        embeddings = get_azure_embeddings(texts)
        
        if using_qdrant:
            # Using Qdrant
            from qdrant_client import models
            
            # Prepare points for Qdrant
            points = []
            for j, (text, metadata, embedding) in enumerate(zip(texts, metadatas, embeddings)):
                point_id = i + j
                points.append(models.PointStruct(
                    id=point_id,
                    vector=embedding,  # Azure embeddings are already in list format
                    payload={
                        "text": text,
                        "metadata": metadata
                    }
                ))
            
            # Upload points to Qdrant
            storage.upsert(
                collection_name=COLLECTION_NAME,
                points=points
            )
        else:
            # Using local storage
            payloads = []
            for text, metadata in zip(texts, metadatas):
                payloads.append({
                    "text": text,
                    "metadata": metadata
                })
            
            storage.upsert(embeddings, payloads)
        
        logger.info(f"Indexed batch {i//batch_size + 1}/{(len(chunks)-1)//batch_size + 1}, {len(batch_chunks)} chunks")

    """Simple local vector storage implementation"""
    def __init__(self, save_path="./vector_store.pkl"):
        self.vectors = []
        self.payloads = []
        self.save_path = save_path
        self._load()
    
    def _load(self):
        """Load existing data if available"""
        if os.path.exists(self.save_path):
            try:
                with open(self.save_path, "rb") as f:
                    data = pickle.load(f)
                    self.vectors = data.get("vectors", [])
                    self.payloads = data.get("payloads", [])
                logger.info(f"Loaded {len(self.vectors)} vectors from {self.save_path}")
            except Exception as e:
                logger.error(f"Error loading vector store: {str(e)}")
    
    def _save(self):
        """Save current data"""
        os.makedirs(os.path.dirname(self.save_path), exist_ok=True)
        with open(self.save_path, "wb") as f:
            pickle.dump({
                "vectors": self.vectors,
                "payloads": self.payloads
            }, f)
        logger.info(f"Saved {len(self.vectors)} vectors to {self.save_path}")
    
    def upsert(self, vectors, payloads):
        """Add vectors and payloads"""
        self.vectors.extend(vectors)
        self.payloads.extend(payloads)
        self._save()
    
    def search(self, query_vector, limit=5):
        """Search for similar vectors"""
        if not self.vectors:
            return []
        
        # Convert query_vector to numpy array
        query_vector = np.array(query_vector)
        
        # Calculate cosine similarity
        similarities = []
        for vector in self.vectors:
            vector = np.array(vector)
            similarity = np.dot(query_vector, vector) / (np.linalg.norm(query_vector) * np.linalg.norm(vector))
            similarities.append(similarity)
        
        # Get top k indices
        top_indices = np.argsort(similarities)[-limit:][::-1]
        
        # Format results
        results = []
        for idx in top_indices:
            results.append({
                "text": self.payloads[idx]["text"],
                "metadata": self.payloads[idx]["metadata"],
                "score": similarities[idx]
            })
        
        return results

# Search function for testing
def search_documents(query, storage, top_k=5, keyword=None):
    """Hybrid search: Combine vector similarity with keyword filtering"""
    # Compute query embedding
    query_embedding = get_azure_embeddings([query])[0]
    
    # Check if storage is from Qdrant or local
    try:
        from qdrant_client import QdrantClient
        using_qdrant = isinstance(storage, QdrantClient)
    except ImportError:
        using_qdrant = False
    
    # Perform vector similarity search
    if using_qdrant:
        # Using Qdrant
        results = storage.search(
            collection_name=COLLECTION_NAME,
            query_vector=query_embedding,
            limit=top_k
        )
        
        # Format results
        formatted_results = []
        for res in results:
            formatted_results.append({
                "text": res.payload["text"],
                "metadata": res.payload["metadata"],
                "score": res.score
            })
    else:
        # Using local storage
        formatted_results = storage.search(query_embedding, limit=top_k)
    
    # Apply keyword filtering if a keyword is provided
    if keyword:
        filtered_results = [
            result for result in formatted_results
            if keyword.lower() in result["text"].lower() or
               any(keyword.lower() in str(value).lower() for value in result["metadata"].values())
        ]
    else:
        filtered_results = formatted_results
    
    return filtered_results

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Document Indexing with Azure OpenAI Embeddings")
    parser.add_argument("--content_dir", type=str, default="content", help="Path to the content directory")
    parser.add_argument("--local", action="store_true", help="Use local vector storage instead of Qdrant")
    parser.add_argument("--search", type=str, help="Search query to test the system")
    parser.add_argument("--keyword", type=str, help="Keyword for hybrid search")
    parser.add_argument("--azure_key", type=str, help="Azure OpenAI API key")
    parser.add_argument("--azure_endpoint", type=str, help="Azure OpenAI endpoint")
    parser.add_argument("--deployment", type=str, help="Azure OpenAI embedding deployment name")
    parser.add_argument("--api_version", type=str, help="Azure OpenAI API version")
    parser.add_argument("--debug", action="store_true", help="Enable debug mode with verbose logging")
    args = parser.parse_args()
    
    # Enable debug logging if requested
    if args.debug:
        logger.setLevel(logging.DEBUG)
        # Add a handler to output to console
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.DEBUG)
        logger.addHandler(console_handler)
        logger.debug("Debug mode enabled")
    
    # Set the content directory path
    content_directory = args.content_dir
    
    # Update Azure configuration if provided
    if args.azure_key:
        AZURE_OPENAI_KEY = args.azure_key
        logger.info("Using Azure OpenAI API key from command line")
    if args.azure_endpoint:
        AZURE_OPENAI_ENDPOINT = args.azure_endpoint
        logger.info(f"Using Azure OpenAI endpoint from command line: {AZURE_OPENAI_ENDPOINT}")
    if args.deployment:
        EMBEDDING_DEPLOYMENT = args.deployment
        logger.info(f"Using Azure OpenAI deployment from command line: {EMBEDDING_DEPLOYMENT}")
    if args.api_version:
        AZURE_OPENAI_API_VERSION = args.api_version
        logger.info(f"Using Azure OpenAI API version from command line: {AZURE_OPENAI_API_VERSION}")
        
    # Validate configuration
    if not AZURE_OPENAI_KEY:
        logger.error("Azure OpenAI API key is required. Set it with --azure_key or AZURE_OPENAI_API_KEY environment variable.")
        exit(1)
    if not AZURE_OPENAI_ENDPOINT:
        logger.error("Azure OpenAI endpoint is required. Set it with --azure_endpoint or AZURE_OPENAI_ENDPOINT environment variable.")
        exit(1)
    if not EMBEDDING_DEPLOYMENT:
        logger.error("Azure OpenAI deployment name is required. Set it with --deployment or AZURE_OPENAI_EMBEDDING_DEPLOYMENT environment variable.")
        exit(1)
        
    # Fix the endpoint URL format if needed
    if AZURE_OPENAI_ENDPOINT.endswith('/'):
        AZURE_OPENAI_ENDPOINT = AZURE_OPENAI_ENDPOINT[:-1]
    if not AZURE_OPENAI_ENDPOINT.startswith('https://'):
        AZURE_OPENAI_ENDPOINT = f"https://{AZURE_OPENAI_ENDPOINT}"
    
    # Initialize storage
    if args.local:
        storage = LocalVectorStorage()
    else:
        storage = try_init_qdrant()
    
    # Run indexing
    if not args.search:
        # Load documents
        documents = load_documents(content_directory)
        
        # Split into chunks
        chunks = split_documents(documents)
        
        # Index in storage
        index_documents(chunks, storage)
        
        logger.info(f"Indexing complete. {len(chunks)} chunks processed.")
    
    # Test search if requested
    if args.search:
        results = search_documents(args.search, storage, keyword=args.keyword)
        
        print("\nSearch Results:")
        print("-" * 80)
        for i, result in enumerate(results):
            print(f"Result {i+1} (Score: {result['score']:.4f}):")
            print(f"Text: {result['text']}")
            print(f"Source: {result['metadata'].get('source', 'Unknown')}")
            print(f"Page Number: {result['metadata'].get('page_label', 'Unknown')}")
            print("-" * 80)